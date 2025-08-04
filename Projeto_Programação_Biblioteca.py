from docx import Document
import os

# -------------------- Registros e listagens de dados --------------------

def registar_livros(livros):
    n = int(input("Quantos livros deseja registrar? "))
    for _ in range(n):
        nome = input("Nome do livro: ")
        autor = input(f"Autor do livro {nome}: ")
        ano = int(input("Ano de publicação: "))
        livros.append(Livro(nome, autor, ano))
        
def registar_utilizadores(utilizadores):
    n = int(input("Quantos utilizadores deseja registrar? "))
    for i in range(n):
        nome = input("Nome do utilizador: ")
        contato = int(input("Contato: "))
        ID = len(utilizadores)
        utilizadores.append(Utilizador(nome, ID, contato))
        
def fazer_emprestimo(emprestimos, livros, utilizadores):
    listar_livros(livros)
    id_l = int(input("Informe o ID do livro: "))
    listar_utilizadores(utilizadores)
    id_u = int(input("Informe o ID do utilizador: "))
    if id_l >= len(livros) or id_u >= len(utilizadores) or not livros[id_l].disponivel:
        print("Erro! ID inválido ou livro indisponível.")
        return
    data = input("Data do empréstimo (dd/mm/aaaa): ")
    previsão_devo = input("Data prevista de devolução (dd/mm/aaaa)")
    livros[id_l].disponivel = False
    emprestimos.append(Emprestimo(id_l, id_u, data, previsão_devo))

def fazer_devolucao(emprestimos, livros, utilizadores):
    id_l = int(input("ID do livro a devolver: "))
    for e in emprestimos:
        if e.id_livro == id_l and e.data_devolucao == "-":
            e.data_devolucao = input("Data da devolução (dd/mm/aaaa): ")
            livros[id_l].disponivel = True
            print("Livro devolvido com sucesso!")
            guardar_livros_docx(livros)
            guardar_emprestimos_docx(emprestimos, livros, utilizadores)
            return
    print("Erro: empréstimo não encontrado ou já devolvido.")
        
def listar_livros(livros):
    print("\n--- Livros ---")
    for i, l in enumerate(livros):
        print(f"ID: {i} | Nome: {l.nome} | Autor: {l.autor} | Ano: {l.ano} | Disponível: {'Sim' if l.disponivel else 'Não'}")
        
def listar_utilizadores(utilizadores):
    print("\n--- Utilizadores ---")
    for i, u in enumerate(utilizadores):
        print(f"ID: {i} | Nome: {u.nome} | Contato: {u.contato}")
        
def listar_emprestimos(emprestimos, livros, utilizadores):
    print("\n--- Empréstimos ---")
    for e in emprestimos:
        if e.id_livro < len(livros) and e.id_utilizador < len(utilizadores):
            livro = livros[e.id_livro].nome
            user = utilizadores[e.id_utilizador].nome
            print(f"Livro: {livro} | Utilizador: {user} | Empréstimo: {e.data_emprestimo} | Devolução prevista: {e.data_prevista_devolucao} | Devolução: {e.data_devolucao if e.data_devolucao != '-' else 'Pendente'}")

# -------------------- Atualização de dados --------------------

def atualizar_dados_livro(livros):
    listar_livros(livros)
    id_l = int(input("ID do livro a atualizar: "))
    if 0 <= id_l < len(livros):
        livros[id_l].nome = input("Novo nome do livro: ")
        livros[id_l].autor = input("Novo autor: ")
        livros[id_l].ano = int(input("Novo ano: "))
        print("Livro atualizado com sucesso!")
    else:
        print("ID inválido.")
        
def atualizar_disponibilidade_livros(livros, emprestimos):
    for livro in livros:
        emprestado = any(
            e.id_livro == livros.index(livro) and e.data_devolucao == "-"
            for e in emprestimos
        )
        livro.disponivel = not emprestado
    print("\nDisponibilidade dos livros atualizada com base nos empréstimos.\n")
        
def atualizar_dados_utilizador(utilizadores):
    listar_utilizadores(utilizadores)
    id_u = int(input("ID do utilizador a atualizar: "))
    if 0 <= id_u < len(utilizadores):
        utilizadores[id_u].nome = input("Novo nome: ")
        utilizadores[id_u].contato = int(input("Novo contato: "))
        print("Utilizador atualizado com sucesso!")
    else:
        print("ID inválido.")
        
def atualizar_dados_emprestimo(emprestimos, livros, utilizadores):
    listar_emprestimos(emprestimos, livros, utilizadores)
    idx = int(input("Número do empréstimo a atualizar (posição na lista): "))
    if 0 <= idx < len(emprestimos):
        id_l = int(input("Novo ID do livro: "))
        id_u = int(input("Novo ID do utilizador: "))
        data = input("Nova data de empréstimo (dd/mm/aaaa): ")
        previsão_devo = input("Nova data prevista de devolução (dd/mm/aaaa): ")
        if id_l < len(livros) and id_u < len(utilizadores):
            emprestimos[idx].id_livro = id_l
            emprestimos[idx].id_utilizador = id_u
            emprestimos[idx].data_emprestimo = data
            emprestimos[idx].data_prevista_devolucao = previsão_devo
            print("Dados do empréstimo atualizados!")
        else:
            print("IDs inválidos.")
    else:
        print("Índice inválido.")
        
def atualizar_dados_devolucao(emprestimos, livros):
    listar_emprestimos(emprestimos, livros, [])
    idx = int(input("Número do empréstimo a atualizar devolução: "))
    if 0 <= idx < len(emprestimos):
        if emprestimos[idx].data_devolucao != "-":
            nova_data = input("Nova data de devolução (dd/mm/aaaa): ")
            emprestimos[idx].data_devolucao = nova_data
            print("Data de devolução atualizada.")
        else:
            print("Este empréstimo ainda não foi devolvido.")
    else:
        print("Índice inválido.")

# -------------------- Guardar e carregar dados --------------------

def guardar_livros_docx(livros):
    doc = Document()
    doc.add_heading("Lista de Livros", 0)
    for i, l in enumerate(livros):
        doc.add_paragraph(
            f"ID: {i} | Nome: {l.nome} | Autor: {l.autor} | Ano: {l.ano} | Disponível: {'Sim' if l.disponivel else 'Não'}"
        )
    doc.save("livros.docx")
        
def guardar_utilizadores_docx(utilizadores):
    doc = Document()
    doc.add_heading("Lista de Utilizadores", 0)
    for u in utilizadores:
        doc.add_paragraph(
            f"ID: {u.ID} | Nome: {u.nome} | Contato: {u.contato}"
        )
    doc.save("utilizadores.docx")
        
def guardar_emprestimos_docx(emprestimos, livros, utilizadores):
    doc = Document()
    doc.add_heading("Lista de Empréstimos", 0)
    for e in emprestimos:
        nome_livro = livros[e.id_livro].nome if e.id_livro < len(livros) else "Desconhecido"
        nome_user = utilizadores[e.id_utilizador].nome if e.id_utilizador < len(utilizadores) else "Desconhecido"
        doc.add_paragraph(
            f"Livro: {nome_livro} | Utilizador: {nome_user} | Empréstimo: {e.data_emprestimo} | Previsão de devolução: {e.data_prevista_devolucao} | Devolução: {e.data_devolucao if e.data_devolucao != '-' else 'Pendente'}"
        )
    doc.save("emprestimos.docx")
        
def carregar_livros():
    livros = []
    if os.path.exists("livros.txt"):
        with open("livros.txt", "r") as f:
            for linha in f:
                nome, autor, ano, disp = linha.strip().split(";")
                livros.append(Livro(nome, autor, int(ano), bool(int(disp))))
    return livros
        
def carregar_livros_docx():
    livros = []
    if os.path.exists("livros.docx"):
        doc = Document("livros.docx")
        for par in doc.paragraphs[1:]:
            partes = par.text.split(" | ")
            nome = partes[1].split(": ")[1]
            autor = partes[2].split(": ")[1]
            ano = int(partes[3].split(": ")[1])
            disponivel = partes[4].split(": ")[1] == "Sim"
            livros.append(Livro(nome, autor, ano, disponivel))
    return livros
        
def carregar_utilizadores_docx():
    utilizadores = []
    if os.path.exists("utilizadores.docx"):
        doc = Document("utilizadores.docx")
        for par in doc.paragraphs[1:]:
            if not par.text.strip():
                continue
            partes = par.text.split(" | ")
            if len(partes) < 3:
                print(f"Parágrafo ignorado (formato inválido): {par.text}")
                continue
            try:
                ID = int(partes[0].split(": ")[1])
                nome = partes[1].split(": ")[1]
                contato = int(partes[2].split(": ")[1])
                utilizadores.append(Utilizador(nome, ID, contato))
            except Exception as e:
                print(f"Erro ao processar parágrafo: {par.text} -> {e}")
    return utilizadores
        
def carregar_emprestimos_docx(livros, utilizadores):
    emprestimos = []
    if os.path.exists("emprestimos.docx"):
        doc = Document("emprestimos.docx")
        for par in doc.paragraphs[1:]:
            if not par.text.strip():
                continue
            partes = par.text.split(" | ")
            if len(partes) < 5:
                print(f"Parágrafo ignorado (formato inválido): {par.text}")
                continue
            try:
                nome_livro = partes[0].split(": ")[1]
                nome_utilizador = partes[1].split(": ")[1]
                data_e = partes[2].split(": ")[1]
                previsao_devo = partes[3].split(": ")[1]
                data_d = partes[4].split(": ")[1]
                
                id_l = next((i for i, l in enumerate(livros) if l.nome == nome_livro), -1)
                id_u = next((i for i, u in enumerate(utilizadores) if u.nome == nome_utilizador), -1)
                if id_l != -1 and id_u != -1:
                    emprestimos.append(Emprestimo(id_l, id_u, data_e, previsao_devo, data_d if data_d != "Pendente" else "-"))
                else:
                    print(f"Livro ou utilizador não encontrado para: {par.text}")
            except Exception as e:
                print(f"Erro ao processar parágrafo: {par.text} -> {e}")
    return emprestimos

    